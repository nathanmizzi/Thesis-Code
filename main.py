# Imports
import math
import random
import traceback
from datetime import datetime
from bs4 import BeautifulSoup as bs, MarkupResemblesLocatorWarning, XMLParsedAsHTMLWarning
from urllib.parse import urljoin
import requests
import pandas as pd
import warnings
import uuid as uid
from matplotlib import pyplot as plt
from docx import Document
from docx.shared import Inches, RGBColor

# Metrics for pdf generation:
total_seconds = {}
successful_hit_type = {}
hit_type = {}
sqliStringsPerWebsite = {}
sqliStringsAttemptedInTotal = {}

safeWebPagesInSite = {}
vulnerableWebPagesInSite = {}

hit_type["DVWA"] = []
hit_type["XVWA"] = []
hit_type["Orange_HRM"] = []
hit_type["Mutillidae"] = []
hit_type["Webgoat"] = []
hit_type["Juice_Shop"] = []
hit_type["Bodgeit"] = []

# A Dictionary containing dataframes which contain info regarding a website
reportDetails = {}

# Contains all the Base URL's of the webpages to test on

# It is important that the names of the sites below match the respective csv files name perfectly.
# list_of_source_csvs = ["BWAPP", "DVWA", "Mutillidae", "Orange_HRM", "Webgoat", "XVWA"]
list_of_source_csvs = ["Bodgeit"]

urls_to_test = {}
vulnerable_urls = []
sqliStrings = []
tested_urls = []

directoryListPath = "utils/directoryLists/dirbuster_200.txt"
passwordListPath = "utils/passwordLists/passlist.txt"
subdomainListPath = "utils/subdomainLists/subdomains.txt"
usernameListPath = "utils/usernameLists/usernames_small.txt"

# Utilities

def request(url):
    try:
        return requests.get(url.trim())
    except requests.exceptions.ConnectionError:
        pass
    except requests.exceptions.InvalidURL:
        print("Failed To Check URL: " + url)

def getCurrentDateTime():
    currentDate = datetime.now()
    return currentDate

def differenceInSeconds(stardDate, endDate):
    difference = (endDate - stardDate)

    return difference.total_seconds()

def generateReports():

    try:

        dateTimeToday = datetime.now().strftime('%d-%m-%Y_%H:%M:%S')

        report = Document()

        h1 = report.add_heading("Penetration Testing Report:", 0)
        h1.bold = True

        if len(hit_type['DVWA']) > 0:
            report.add_paragraph(f"\nAll Types of attacks attempted: {hit_type['DVWA']}\n")

        report.add_paragraph("General Scan information: \n")

        generalTable = report.add_table(rows=len(list_of_source_csvs), cols=4)

        row = generalTable.rows[0].cells
        row[0].text = 'Website'
        row[1].text = 'No. Of URL\'s Tested'
        row[2].text = 'No. Of Successful Attacks'
        row[3].text = 'Time Taken'

        for website in list_of_source_csvs:
            row = generalTable.add_row().cells
            row[0].text = website
            row[1].text = str((len(safeWebPagesInSite[website]) + len(vulnerableWebPagesInSite[website])))
            row[2].text = str(len(vulnerableWebPagesInSite[website]))
            row[3].text = str(math.ceil(total_seconds[website])) + " Seconds"

        for idx, col in enumerate(generalTable.columns): col.width = Inches(1.3)

        generalTable.style = 'Colorful List'

        report.add_paragraph("\nWebsite Vulnerability Statuses: \n")

        imagesDir = "./assets"

        safetyTable = report.add_table(rows=len(list_of_source_csvs), cols=2)

        row = safetyTable.rows[0].cells
        row[0].text = 'Website'
        row[1].text = 'Safety'

        for website in list_of_source_csvs:
            row = safetyTable.add_row().cells
            row[0].text = website

            img_paragraph = row[1].paragraphs[0]
            if len(vulnerableWebPagesInSite[website]) <= 0:
                img_paragraph.text = "Website Secure! "
                img_paragraph.add_run().add_picture(f"{imagesDir}/Check.png", width=Inches(0.4), height=Inches(0.3))
            else:
                img_paragraph.text = "Website Vulnerable! "
                img_paragraph.add_run().add_picture(f"{imagesDir}/Cross.png", width=Inches(0.4), height=Inches(0.3))

        for idx, col in enumerate(safetyTable.columns): col.width = Inches(1.3)

        safetyTable.style = 'Colorful List'

        report.add_page_break()

        for website in list_of_source_csvs:

            websiteHeader = report.add_heading(f"{website}", 0)
            websiteHeader.bold = True

            report.add_paragraph(f"Date and Time of Report: {datetime.now().strftime('%d-%m-%Y, %H:%M:%S')}")

            if len(successful_hit_type[website]) > 0:
                report.add_paragraph(f"\nTypes of Successful attacks: {successful_hit_type[website]}")

                report.add_paragraph(f"\nSQLI Strings Attempted in Total: {sqliStringsAttemptedInTotal[website]}\n")

                sqliStringHeader = report.add_heading('SQLI Strings Used: \n', 3)
                sqliStringHeader.style.font.color.rgb = RGBColor(0, 0, 0)

                for injectionString in sqliStringsPerWebsite[website]:
                    report.add_paragraph("          [-]   " + injectionString)

                vulnerableLinkHeader = report.add_heading('Vulnerable Pages: \n', 3)
                vulnerableLinkHeader.style.font.color.rgb = RGBColor(0, 0, 0)

                for url in vulnerableWebPagesInSite[website]:
                    report.add_paragraph("          [-]   " + url)

            report.add_paragraph(f"\nVulnerability Statistics: ")

            # Creating a Pie Chart of vulnerable and safe code.
            lbls = ["Vulnerable", "Safe"]
            id_of_img = uid.uuid1().__str__()
            fig = plt.figure(figsize= (3, 2))
            plt.pie([len(vulnerableWebPagesInSite[website]), len(safeWebPagesInSite[website])], labels=lbls,
                    autopct='%1.1f%%', pctdistance=0.85, shadow=True, explode=(0.05, 0.05), startangle=90)

            donut = plt.Circle((0, 0), 0.70, fc='white')
            fig = plt.gcf()

            fig.gca().add_artist(donut)
            plt.title("URL Vulnerability Distribution")

            plt.savefig('utils/graphs/' + id_of_img + ".png")
            report.add_picture('utils/graphs/' + id_of_img + ".png", width=Inches(3), height=Inches(2))

            report.add_paragraph(f"\nRecommendations: ")

            if len(successful_hit_type[website]) > 0:
                sqliAdded = False
                for string in successful_hit_type[website]:
                    if "SQL Injection" in string and sqliAdded is False:
                        recommendedPar = report.add_paragraph("SQL Injection: https://owasp.org/www-community/attacks/SQL_Injection")
                        recommendedParFormat = recommendedPar.paragraph_format
                        recommendedParFormat.left_indent = Inches(0.5)
                        sqliAdded = True
            else:
                recommendedPar = report.add_paragraph("Nothing To Recommend, Your Webpage is Secure!")
                recommendedParFormat = recommendedPar.paragraph_format
                recommendedParFormat.left_indent = Inches(0.5)

            report.add_page_break()

            report.save(f"reports/report_{dateTimeToday}.docx")

    except Exception as e:
        print("\nError In Report Generation! :")
        print(e)
        traceback.print_exc()


# This function simulates a penetration test on a website, and allows the reports to be generated without scrapes
def populateTestData(websiteName, doSafe):

    # Random Values to choose from
    typesOfAttacksDummyData = ["Error Based SQL Injection"]
    sqliStringsDummyData = ["=<1 or '1'>", "1 or 1'", 'select * from users where id  =  1 or 1#""?  =  1 or 1  =  1 -- 1",1,,']
    total_secondsDummyData = [50, 110, 27, 43]
    vulnerableWebPagesInSiteDummyData = ["http://" + websiteName + "/login.php", "http://" + websiteName + "/sqli.php", "http://" + websiteName + "/bruteForce.php"]

    # Instantiating The Necessary Variables
    total_seconds[websiteName] = 0
    sqliStringsAttemptedInTotal[websiteName] = 0
    successful_hit_type[websiteName] = []
    vulnerableWebPagesInSite[websiteName] = []
    safeWebPagesInSite[websiteName] = []
    sqliStringsPerWebsite[websiteName] = []

    # Populating the variables as a normal scrape would
    if not doSafe:
        hit_type[websiteName].append(typesOfAttacksDummyData[0])
        successful_hit_type[websiteName].append(typesOfAttacksDummyData[0])
        sqliStringsAttemptedInTotal[websiteName] += random.randint(0, 100)

    for i in range(len(sqliStringsDummyData) - 1):
        valueSelected = random.randint(0, len(sqliStringsDummyData) - 1)
        sqliStringsPerWebsite[websiteName].append(sqliStringsDummyData[valueSelected])

    for i in range(len(total_secondsDummyData) - 1):
        valueSelected = random.randint(0, len(total_secondsDummyData) - 1)
        total_seconds[websiteName] = total_secondsDummyData[valueSelected]

    for i in range(len(vulnerableWebPagesInSiteDummyData) - 1):
        if not doSafe:
            valueSelected = random.randint(0, len(vulnerableWebPagesInSiteDummyData) - 1)
            vulnerable_urls.append(vulnerableWebPagesInSiteDummyData[valueSelected])
            vulnerableWebPagesInSite[websiteName].append(vulnerableWebPagesInSiteDummyData[valueSelected])

    ratioOfSafeSites = total_secondsDummyData[random.randint(0, len(total_secondsDummyData) - 1)]

    for i in range(ratioOfSafeSites):
        safeWebPagesInSite[websiteName].append("http://" + websiteName + "/safeSite.php")

def readWebsiteLinksCsv():

    for fileName in list_of_source_csvs:

        websites_urls = []

        df = pd.read_csv('utils/websiteLinks/' + fileName + '.csv')

        for index, row in df.iterrows():
            try:
                if df.iloc[index]["Processed"] == True and df.iloc[index]["Method"] == "GET":
                    # print(df.iloc[index]["URI"])
                    websites_urls.append(df.iloc[index]["URI"])
            except Exception as e:
                print(e)

        urls_to_test[fileName] = websites_urls

def readSQLICsv():

    df = pd.read_csv('utils/SQLI_Dataset/SQLIV3_Shrunken.csv')

    for index, row in df.iterrows():
        try:
            sqliStrings.append(df.iloc[index]["Sentence"])
        except Exception as e:
            print(e)

# Brute Forcers
def websiteBruteForce(page_url):
    with open(usernameListPath, "r") as usernames:

        for username in usernames:

            with open(passwordListPath, "r") as passwords:

                for password in passwords:

                    username = username.strip()
                    password = password.strip()
                    print(f"[!!] Trying To Brute Force With Username: '{username}', Password: '{password}'")

                    with requests.Session() as s:

                        resp = s.get(page_url)
                        parsed_html = bs(resp.content, features="html.parser")
                        input_value = parsed_html.body.find('input', attrs={'name': 'user_token'}).get("value")
                        data_dict = {"username": username, "password": password, "Login": "Login",
                                     "user_token": input_value}
                        response = s.post(page_url, data_dict)

                    if b"Login failed" in response.content:
                        pass
                    else:
                        print("\n[+] Username: --> " + username)
                        print("[+] Password: --> " + password)

                        return True

        print("\nUsername and password are not in wordlists.")
        return False


# Error Based SQL Injection

def get_all_forms(url, cookies):

    if cookies is not None:
        response = s.get(url, cookies=cookies)
        response_content = response.content

        soup = bs(response_content.decode('utf-8', 'ignore'), "html.parser")

        return soup.find_all("form")
    else:
        response = s.get(url)
        response_content = response.content

        soup = bs(response_content.decode('utf-8', 'ignore'), "html.parser")

        return soup.find_all("form")

def get_form_details(form):
    details = {}

    try:
        action = form.attrs.get("action").lower()
    except:
        action = None

    method = form.attrs.get("method", "get").lower()

    inputs = []

    for input_tag in form.find_all("input"):
        input_type = input_tag.attrs.get("type", "text")
        input_name = input_tag.attrs.get("name")
        input_value = input_tag.attrs.get("value", "")
        inputs.append({"type": input_type, "name": input_name, "value": input_value})

    details["action"] = action
    details["method"] = method
    details["inputs"] = inputs

    # print(details)
    return details

def isInjectable(response):
    errors = {
        "you have an error in your sql syntax;",
        "warning: mysql",
        "unclosed quotation mark after the character string",
        "quoted string not properly terminated",
        "Uncaught mysqli_sql_exception",
        "unexpected token"
    }

    for error in errors:
        if response.status_code != 404:
            if error in response.content.decode(errors='ignore').lower():
                return True

    return False

def sqlInjectionScan(url, cookies, nameOfWebsite):

    for sqliString in sqliStrings:

        sqliStringsAttemptedInTotal[nameOfWebsite] += 1

        if url not in tested_urls:

            print("TESTING URL: " + url)

            forms = get_all_forms(url, cookies)

            if len(forms) >= 1:
                print(f"[+] Detected {len(forms)} forms on {url}.")
            else:
                tested_urls.append(url)
                return False

            for form in forms:
                form_details = get_form_details(form)

                try:
                    if "xvwa" in url and form_details["inputs"][0]["name"] == "username": # Adding condition to skip login form for XVWA as it was causing issues
                        continue
                except:
                    continue

                for c in "\"'":

                    data = {}
                    for input_tag in form_details["inputs"]:
                        if input_tag["type"] == "hidden" or input_tag["value"]:
                            try:
                                data[input_tag["name"]] = input_tag["value"] + c
                            except:
                                pass
                        elif input_tag["type"] != "submit":
                              # data[input_tag["name"]] = f"<1'1 or 1>" # New code that aims to catch more errors on a multitude of websites
                              data[input_tag["name"]] = sqliString  # Takes a string from the Dataset containing sqli strings.

                    url = urljoin(url, form_details["action"])
                    if form_details["method"] == "post":
                        s.cookies.clear()
                        res = s.post(url, data=data)
                    elif form_details["method"] == "get":
                        s.cookies.clear()
                        res = s.get(url, params=data)

                    if isInjectable(res):
                        print("[+] SQL Injection vulnerability detected, link:", url)
                        print("[+] Form:")
                        print(form_details)

                        sqliStringsPerWebsite[nameOfWebsite].append(sqliString)
                        tested_urls.append(url)

                        return True

    tested_urls.append(url)
    return False


# BLind SQL Injection

def blind_sql(url, cookies):

    print("TESTING URL: " + url)

    try:

        setSleepTime = 2

        # Steps

        # Get Form Details
        forms = get_all_forms(url, cookies)

        if len(forms) >= 1:
            print(f"[+] Detected {len(forms)} forms on {url}.")

        for form in forms:
            form_details = get_form_details(form)

            try:
                if "xvwa" in url and form_details["inputs"][0]["name"] == "username":  # Adding condition to skip login form for XVWA as it was causing issues
                    continue
            except:
                continue

            for c in "\"'":

                data = {}
                for input_tag in form_details["inputs"]:
                    if input_tag["type"] == "hidden" or input_tag["value"]:
                        try:
                            data[input_tag["name"]] = input_tag["value"] + c
                        except:
                            pass
                    elif input_tag["type"] != "submit":
                        # Inject SQL and sleep command
                        data[input_tag["name"]] = f"SLEEP( " + str(setSleepTime) +")/*' || SLEEP(" + str(setSleepTime) + ") || '\" || SLEEP(" + str(setSleepTime) + ") || \"*/"

                url = urljoin(url, form_details["action"])

                # Get Response
                if form_details["method"] == "post":
                    s.cookies.clear()
                    res = s.post(url, data=data)
                elif form_details["method"] == "get":
                    s.cookies.clear()
                    res = s.get(url, params=data)

                requestTime = res.elapsed.total_seconds()
                tested_urls.append(url)

                # If time to respond is greater than the sleep command, assume that blind SQL Injection is possible
                if requestTime >= setSleepTime:

                    print("[+] BLind SQL Injection vulnerability detected, link:", url)
                    print("[+] Form:")
                    print(form_details)

                    return True
                # Else, command ignored, SQL Injection failed
                else:
                    pass

        return False

    except:
        print("Error analysing URL: " + url)
        # traceback.print_exc()
        return False

# Website authenticators

def DVWA_error_based(urls):

    print("\n --- Error Based SQL --- \n")

    total_seconds["DVWA"] = 0
    timeStarted = getCurrentDateTime()
    sqliStringsPerWebsite["DVWA"] = []

    try:
        sqliStringsAttemptedInTotal["DVWA"] = 0
        successful_hit_type["DVWA"] = []
        vulnerableWebPagesInSite["DVWA"] = []
        safeWebPagesInSite["DVWA"] = []
        hit_type["DVWA"].append("Error Based SQL Injection")

        # Firstly, create a logged-in session in order to create requests

        s.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64; rv:46.0) Gecko/20100101 Firefox/46.0',
            'Cookie': 'security=low; PHPSESSID=geo7gb3ehf5gfnbhrvuqu545i7'
        }

        resp = s.get('http://127.0.0.1/login.php')
        parsed_html = bs(resp.content, features="html.parser")
        input_value = parsed_html.body.find('input', attrs={'name': 'user_token'}).get("value")
        data_dict = {"username": 'admin', "password": 'password', "Login": "Login",
                     "user_token": input_value}

        response = s.post('http://127.0.0.1/login.php', data_dict)
        responseContent = response.content
        cookies = response.cookies

        # Then test other urls

        for url in urls["DVWA"]:

            if sqlInjectionScan(url, cookies, "DVWA"):

                vulnerableWebPagesInSite["DVWA"].append(url)

                vulnerable_urls.append("DVWA: " + url)

                if "Error Based SQL Injection" not in successful_hit_type["DVWA"]:
                    successful_hit_type["DVWA"].append("Error Based SQL Injection")
            else:
                safeWebPagesInSite["DVWA"].append(url)
    except Exception as e:
        print("\nDVWA Error: \n")
        print(e)
        traceback.print_exc()

    # Get the current time when all tests end
    timeEnded = getCurrentDateTime()
    total_seconds["DVWA"] += differenceInSeconds(timeStarted, timeEnded)

def DVWA_Blind(urls):

    print("\n --- BLIND SQL --- \n")

    total_seconds["DVWA"] = 0
    timeStarted = getCurrentDateTime()
    sqliStringsPerWebsite["DVWA"] = []

    try:

        sqliStringsAttemptedInTotal["DVWA"] = 0
        successful_hit_type["DVWA"] = []
        vulnerableWebPagesInSite["DVWA"] = []
        safeWebPagesInSite["DVWA"] = []
        hit_type["DVWA"].append("BLind SQL Injection")

        # Firstly, create a logged-in session in order to create requests

        s.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64; rv:46.0) Gecko/20100101 Firefox/46.0',
            'Cookie': 'security=low; PHPSESSID=geo7gb3ehf5gfnbhrvuqu545i7'
        }

        resp = s.get('http://127.0.0.1/login.php')
        parsed_html = bs(resp.content, features="html.parser")
        input_value = parsed_html.body.find('input', attrs={'name': 'user_token'}).get("value")
        data_dict = {"username": 'admin', "password": 'password', "Login": "Login",
                     "user_token": input_value}

        response = s.post('http://127.0.0.1/login.php', data_dict)
        responseContent = response.content
        cookies = response.cookies


        # Then test other urls

        for url in urls["DVWA"]:

            if blind_sql(url, cookies):

                vulnerableWebPagesInSite["DVWA"].append(url)

                vulnerable_urls.append("DVWA: " + url)

                if "Blind SQL Injection" not in successful_hit_type["DVWA"]:
                    successful_hit_type["DVWA"].append("Blind SQL Injection")
            else:
                safeWebPagesInSite["DVWA"].append(url)

        timeEnded = getCurrentDateTime()
        total_seconds["DVWA"] += differenceInSeconds(timeStarted, timeEnded)

    except Exception as e:
        print("\nBlind DVWA Error: \n")
        print(e)
        traceback.print_exc()

def XVWA_error_based(urls):

    total_seconds["XVWA"] = 0
    timeStarted = getCurrentDateTime()
    sqliStringsPerWebsite["XVWA"] = []

    try:
        sqliStringsAttemptedInTotal["XVWA"] = 0
        successful_hit_type["XVWA"] = []
        vulnerableWebPagesInSite["XVWA"] = []
        safeWebPagesInSite["XVWA"] = []
        hit_type["XVWA"].append("Error Based SQL Injection")

        for url in urls["XVWA"]:

            if sqlInjectionScan(url, None, "XVWA"):
                vulnerableWebPagesInSite["XVWA"].append(url)

                if url not in vulnerable_urls:
                    vulnerable_urls.append("XVWA: " + url)

                if "Error Based SQL Injection" not in successful_hit_type["XVWA"]:
                    successful_hit_type["XVWA"].append("Error Based SQL Injection")

            else:
                safeWebPagesInSite["XVWA"].append(url)

    except Exception as e:
        print("\nXVWA Error: \n")
        print(e)
        traceback.print_exc()

    timeEnded = getCurrentDateTime()
    total_seconds["XVWA"] += differenceInSeconds(timeStarted, timeEnded)

def XVWA_Blind(urls):

    print("\n --- BLIND SQL --- \n")

    total_seconds["XVWA"] = 0
    timeStarted = getCurrentDateTime()
    sqliStringsPerWebsite["XVWA"] = []

    try:

        sqliStringsAttemptedInTotal["XVWA"] = 0
        successful_hit_type["XVWA"] = []
        vulnerableWebPagesInSite["XVWA"] = []
        safeWebPagesInSite["XVWA"] = []
        hit_type["XVWA"].append("BLind SQL Injection")

        # Then test other urls

        for url in urls["XVWA"]:

            if blind_sql(url, None):

                vulnerableWebPagesInSite["XVWA"].append(url)

                vulnerable_urls.append("XVWA: " + url)

                if "Blind SQL Injection" not in successful_hit_type["XVWA"]:
                    successful_hit_type["XVWA"].append("Blind SQL Injection")
            else:
                safeWebPagesInSite["XVWA"].append(url)

        timeEnded = getCurrentDateTime()
        total_seconds["XVWA"] += differenceInSeconds(timeStarted, timeEnded)

    except Exception as e:
        print("\nBlind XVWA Error: \n")
        print(e)
        traceback.print_exc()

def OrangeHRM_error_based(urls):

    total_seconds["Orange_HRM"] = 0
    timeStarted = getCurrentDateTime()
    sqliStringsPerWebsite["Orange_HRM"] = []

    try:

        successful_hit_type["Orange_HRM"] = []

        # Firstly, create a logged-in session in order to create requests
        resp = s.get('http://localhost:1234/OrangeHRM/symfony/web/index.php/auth/login')
        parsed_html = bs(resp.content, features="html.parser")
        input_value = parsed_html.body.find('input', attrs={'name': '_csrf_token'}).get("value")
        data_dict = {"txtUsername": 'admin', "txtPassword": 'Mcast1234!', "Submit": "LOGIN", "_csrf_token":input_value}

        response = s.post("http://localhost:1234/OrangeHRM/symfony/web/index.php/auth/validateCredentials", data_dict)
        cookies = response.cookies
        response_content = response.content.decode()

        # Then test other urls

        for url in urls["Orange_HRM"]:

            try:
                if sqlInjectionScan(url, cookies, "Orange_HRM"):

                    if url not in vulnerable_urls:
                        vulnerable_urls.append("Orange_HRM: " + url)

                    if "SQL Injection" not in successful_hit_type["Orange_HRM"]:
                        successful_hit_type["Orange_HRM"].append("SQL Injection")
            except:
                continue

    except Exception as e:
        print("\nOrange_HRM Error: \n")
        print(e)
        traceback.print_exc()

    timeEnded = getCurrentDateTime()
    total_seconds["Orange_HRM"] += differenceInSeconds(timeStarted, timeEnded)

def OrangeHRM_Blind(urls):

    print("\n --- BLIND SQL --- \n")

    total_seconds["Orange_HRM"] = 0
    timeStarted = getCurrentDateTime()
    sqliStringsPerWebsite["Orange_HRM"] = []

    try:

        sqliStringsAttemptedInTotal["Orange_HRM"] = 0
        successful_hit_type["Orange_HRM"] = []
        vulnerableWebPagesInSite["Orange_HRM"] = []
        safeWebPagesInSite["Orange_HRM"] = []
        hit_type["Orange_HRM"].append("BLind SQL Injection")

        # Firstly, create a logged-in session in order to create requests

        s.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64; rv:46.0) Gecko/20100101 Firefox/46.0',
            'Cookie': 'security=low; PHPSESSID=geo7gb3ehf5gfnbhrvuqu545i7'
        }

        # Firstly, create a logged-in session in order to create requests
        resp = s.get('http://localhost:1234/OrangeHRM/symfony/web/index.php/auth/login')
        parsed_html = bs(resp.content, features="html.parser")
        input_value = parsed_html.body.find('input', attrs={'name': '_csrf_token'}).get("value")
        data_dict = {"txtUsername": 'admin', "txtPassword": 'Mcast1234!', "Submit": "LOGIN", "_csrf_token":input_value}

        response = s.post("http://localhost:1234/OrangeHRM/symfony/web/index.php/auth/validateCredentials", data_dict)
        cookies = response.cookies
        response_content = response.content.decode()


        # Then test other urls

        for url in urls["Orange_HRM"]:

            if blind_sql(url, cookies):

                vulnerableWebPagesInSite["Orange_HRM"].append(url)

                vulnerable_urls.append("Orange_HRM: " + url)

                if "Blind SQL Injection" not in successful_hit_type["Orange_HRM"]:
                    successful_hit_type["Orange_HRM"].append("Blind SQL Injection")
            else:
                safeWebPagesInSite["Orange_HRM"].append(url)

        timeEnded = getCurrentDateTime()
        total_seconds["Orange_HRM"] += differenceInSeconds(timeStarted, timeEnded)

    except Exception as e:
        print("\nBlind Orange_HRM Error: \n")
        print(e)
        traceback.print_exc()

def Mutillidae_error_based(urls):

    total_seconds["Mutillidae"] = 0
    timeStarted = getCurrentDateTime()
    sqliStringsPerWebsite["Mutillidae"] = []

    try:

        successful_hit_type["Mutillidae"] = []

        for url in urls["Mutillidae"]:

            if sqlInjectionScan(url, None, "Mutillidae"):

                if url not in vulnerable_urls:
                    vulnerable_urls.append("Mutillidae: " + url)

                if "SQL Injection" not in successful_hit_type["Mutillidae"]:
                    successful_hit_type["Mutillidae"].append("SQL Injection")

    except Exception as e:
        print("\nMutillidae Error: \n")
        print(e)
        traceback.print_exc()

    timeEnded = getCurrentDateTime()
    total_seconds["Mutillidae"] += differenceInSeconds(timeStarted, timeEnded)

def Mutillidae_Blind(urls):
    print("\n --- BLIND SQL --- \n")

    total_seconds["Mutillidae"] = 0
    timeStarted = getCurrentDateTime()
    sqliStringsPerWebsite["Mutillidae"] = []

    try:

        sqliStringsAttemptedInTotal["Mutillidae"] = 0
        successful_hit_type["Mutillidae"] = []
        vulnerableWebPagesInSite["Mutillidae"] = []
        safeWebPagesInSite["Mutillidae"] = []
        hit_type["Mutillidae"].append("BLind SQL Injection")

        for url in urls["Mutillidae"]:

            if blind_sql(url, None):

                vulnerableWebPagesInSite["Mutillidae"].append(url)

                vulnerable_urls.append("Mutillidae: " + url)

                if "Blind SQL Injection" not in successful_hit_type["Mutillidae"]:
                    successful_hit_type["Mutillidae"].append("Blind SQL Injection")
            else:
                safeWebPagesInSite["Mutillidae"].append(url)

        timeEnded = getCurrentDateTime()
        total_seconds["Mutillidae"] += differenceInSeconds(timeStarted, timeEnded)

    except Exception as e:
        print("\nBlind Mutillidae Error: \n")
        print(e)
        traceback.print_exc()

def WebGoat_error_based(urls):

    total_seconds["WebGoat"] = 0
    timeStarted = getCurrentDateTime()
    sqliStringsPerWebsite["WebGoat"] = []

    try:

        successful_hit_type["Webgoat"] = []

        # Firstly, create a logged-in session in order to create requests
        s.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64; rv:46.0) Gecko/20100101 Firefox/46.0'
        }

        resp = s.get('http://localhost:8080/WebGoat/login')
        parsed_html = bs(resp.content, features="html.parser")
        data_dict = {"username": 'nathanmizzi', "password": 'mcast1234'}

        response = s.post("http://localhost:8080/WebGoat/login", data_dict)

        response_content = response.content.decode()
        cookies = s.cookies.get_dict()

        # Then test other urls

        for url in urls["Webgoat"]:

            if sqlInjectionScan(url, cookies, "Webgoat"):

                if url not in vulnerable_urls:
                    vulnerable_urls.append("Webgoat: " + url)

                if "SQL Injection" not in successful_hit_type["Webgoat"]:
                    successful_hit_type["Webgoat"].append("SQL Injection")

    except Exception as e:
        print("\nWebgoat Error: \n")
        print(e)
        traceback.print_exc()

    timeEnded = getCurrentDateTime()
    total_seconds["WebGoat"] += differenceInSeconds(timeStarted, timeEnded)

def WebGoat_Blind(urls):

    print("\n --- BLIND SQL --- \n")

    total_seconds["Webgoat"] = 0
    timeStarted = getCurrentDateTime()
    sqliStringsPerWebsite["Webgoat"] = []

    try:

        sqliStringsAttemptedInTotal["Webgoat"] = 0
        successful_hit_type["Webgoat"] = []
        vulnerableWebPagesInSite["Webgoat"] = []
        safeWebPagesInSite["Webgoat"] = []
        hit_type["Webgoat"].append("BLind SQL Injection")

        for url in urls["Webgoat"]:

            if blind_sql(url, None):

                vulnerableWebPagesInSite["Webgoat"].append(url)

                vulnerable_urls.append("Webgoat: " + url)

                if "Blind SQL Injection" not in successful_hit_type["Webgoat"]:
                    successful_hit_type["Webgoat"].append("Blind SQL Injection")
            else:
                safeWebPagesInSite["Webgoat"].append(url)

        timeEnded = getCurrentDateTime()
        total_seconds["Webgoat"] += differenceInSeconds(timeStarted, timeEnded)

    except Exception as e:
        print("\nBlind Webgoat Error: \n")
        print(e)
        traceback.print_exc()

def Juice_Shop_error_based(urls):

    total_seconds["Juice_Shop"] = 0
    timeStarted = getCurrentDateTime()
    sqliStringsPerWebsite["Juice_Shop"] = []

    try:

        sqliStringsAttemptedInTotal["Juice_Shop"] = 0
        successful_hit_type["Juice_Shop"] = []
        vulnerableWebPagesInSite["Juice_Shop"] = []
        safeWebPagesInSite["Juice_Shop"] = []
        hit_type["Juice_Shop"].append("Error Based SQL Injection")

        # Firstly, create a logged-in session in order to create requests
        resp = s.get('http://localhost:3000/#/login')
        parsed_html = bs(resp.content, features="html.parser")
        data_dict = {"email": 'test@email.com', "password": 'password!', "Submit": "loginButton"}

        response = s.post("http://localhost:3000/#/login", data_dict)
        cookies = response.cookies
        response_content = response.content.decode()

        # Then test other urls

        for url in urls["Juice_Shop"]:

            try:

                if sqlInjectionScan(url, cookies, "Juice_Shop"):

                    if url not in vulnerable_urls:
                        vulnerable_urls.append("Juice_Shop: " + url)

                    vulnerableWebPagesInSite["Juice_Shop"].append(url)

                    if "SQL Injection" not in successful_hit_type["Juice_Shop"]:
                        successful_hit_type["Juice_Shop"].append("SQL Injection")

                else:
                    safeWebPagesInSite["Juice_Shop"].append(url)

            except:
                continue

    except Exception as e:
        print("\nJuice_Shop Error: \n")
        print(e)
        traceback.print_exc()

    timeEnded = getCurrentDateTime()
    total_seconds["Juice_Shop"] += differenceInSeconds(timeStarted, timeEnded)

def Juice_Shop_Blind(urls):

    print("\n --- BLIND SQL --- \n")

    total_seconds["Juice_Shop"] = 0
    timeStarted = getCurrentDateTime()
    sqliStringsPerWebsite["Juice_Shop"] = []

    try:

        sqliStringsAttemptedInTotal["Juice_Shop"] = 0
        successful_hit_type["Juice_Shop"] = []
        vulnerableWebPagesInSite["Juice_Shop"] = []
        safeWebPagesInSite["Juice_Shop"] = []
        hit_type["Juice_Shop"].append("BLind SQL Injection")

        # Firstly, create a logged-in session in order to create requests
        resp = s.get('http://localhost:3000/#/login')
        parsed_html = bs(resp.content, features="html.parser")
        data_dict = {"email": 'test@email.com', "password": 'password!', "Submit": "loginButton"}

        response = s.post("http://localhost:3000/#/login", data_dict)
        cookies = response.cookies
        response_content = response.content.decode()

        for url in urls["Juice_Shop"]:

            if blind_sql(url, cookies):

                vulnerableWebPagesInSite["Juice_Shop"].append(url)

                vulnerable_urls.append("Juice_Shop: " + url)

                if "Blind SQL Injection" not in successful_hit_type["Juice_Shop"]:
                    successful_hit_type["Juice_Shop"].append("Blind SQL Injection")
            else:
                safeWebPagesInSite["Juice_Shop"].append(url)

        timeEnded = getCurrentDateTime()
        total_seconds["Juice_Shop"] += differenceInSeconds(timeStarted, timeEnded)

    except Exception as e:
        print("\nBlind Juice_Shop Error: \n")
        print(e)
        traceback.print_exc()

def Bodgeit_error_based(urls):

    print("\n --- Error Based SQL --- \n")

    total_seconds["Bodgeit"] = 0
    timeStarted = getCurrentDateTime()
    sqliStringsPerWebsite["Bodgeit"] = []

    try:
        sqliStringsAttemptedInTotal["Bodgeit"] = 0
        successful_hit_type["Bodgeit"] = []
        vulnerableWebPagesInSite["Bodgeit"] = []
        safeWebPagesInSite["Bodgeit"] = []
        hit_type["Bodgeit"].append("Error Based SQL Injection")


        # Then test other urls

        for url in urls["Bodgeit"]:

            if sqlInjectionScan(url, None, "Bodgeit"):

                vulnerableWebPagesInSite["Bodgeit"].append(url)

                vulnerable_urls.append("Bodgeit: " + url)

                if "Error Based SQL Injection" not in successful_hit_type["Bodgeit"]:
                    successful_hit_type["Bodgeit"].append("Error Based SQL Injection")
            else:
                safeWebPagesInSite["Bodgeit"].append(url)
    except Exception as e:
        print("\nBodgeit Error: \n")
        print(e)
        traceback.print_exc()

    # Get the current time when all tests end
    timeEnded = getCurrentDateTime()
    total_seconds["Bodgeit"] += differenceInSeconds(timeStarted, timeEnded)

def Bodgeit_Blind(urls):
    print("\n --- BLIND SQL --- \n")

    total_seconds["Bodgeit"] = 0
    timeStarted = getCurrentDateTime()
    sqliStringsPerWebsite["Bodgeit"] = []

    try:

        sqliStringsAttemptedInTotal["Bodgeit"] = 0
        successful_hit_type["Bodgeit"] = []
        vulnerableWebPagesInSite["Bodgeit"] = []
        safeWebPagesInSite["Bodgeit"] = []
        hit_type["Bodgeit"].append("BLind SQL Injection")

        for url in urls["Bodgeit"]:

            if blind_sql(url, None):

                vulnerableWebPagesInSite["Bodgeit"].append(url)
                vulnerable_urls.append("Bodgeit: " + url)

                if "Blind SQL Injection" not in successful_hit_type["Bodgeit"]:
                    successful_hit_type["Bodgeit"].append("Blind SQL Injection")
            else:
                safeWebPagesInSite["Bodgeit"].append(url)

        timeEnded = getCurrentDateTime()
        total_seconds["Bodgeit"] += differenceInSeconds(timeStarted, timeEnded)

    except Exception as e:
        print("\nBlind Bodgeit Error: \n")
        print(e)
        traceback.print_exc()

if __name__ == '__main__':

    warnings.filterwarnings(action="ignore", category=MarkupResemblesLocatorWarning)
    warnings.filterwarnings(action="ignore", category=XMLParsedAsHTMLWarning)

    # Populate array of urls from csv files
    readWebsiteLinksCsv()

    # Populate array of SQL Injection strings from Kaggle Dataset
    readSQLICsv()

    testMode = False

    if not testMode:

        if "DVWA" in list_of_source_csvs:
            # Test urls accordingly
            with requests.Session() as s:
                # DVWA_error_based(urls_to_test)
                DVWA_Blind(urls_to_test)

        if "XVWA" in list_of_source_csvs:
            with requests.Session() as s:
                # XVWA_error_based(urls_to_test)
                XVWA_Blind(urls_to_test)

        if "Orange_HRM" in list_of_source_csvs:
            with requests.Session() as s:
            #     OrangeHRM_error_based(urls_to_test)
                OrangeHRM_Blind(urls_to_test)

        # TODO: Investigate Lack of investigated links and failed detections in Mutillidae
        if "Mutillidae" in list_of_source_csvs:
            with requests.Session() as s:
            #     Mutillidae_error_based(urls_to_test)
                Mutillidae_Blind(urls_to_test)

        # TODO: Investigate Lack of detections in the Webgoat urls
        if "Webgoat" in list_of_source_csvs:
            with requests.Session() as s:
            #     WebGoat_error_based(urls_to_test)
                WebGoat_Blind(urls_to_test)

        if "Juice_Shop" in list_of_source_csvs:
            with requests.Session() as s:
                # Juice_Shop_error_based(urls_to_test)
                Juice_Shop_Blind(urls_to_test)

        # TODO: Implement Moodle

        # TODO: Implement bWAPP

        # TODO: Investigate Lack of investigated links and failed detections in Bodgeit
        if "Bodgeit" in list_of_source_csvs:
            with requests.Session() as s:
                Bodgeit_error_based(urls_to_test)
                Bodgeit_Blind(urls_to_test)

        # TODO: Implement WackoPicko

    else:
        populateTestData("DVWA", False)
        populateTestData("XVWA", True)

    # Generate Reports
    generateReports()

    # Urls have been tested, now output the results to the user.
    print("\nVulnerabilities were found with: " + str(len(vulnerable_urls)) + " urls.")
    # print("Total Time Taken To Perform Tests: " + str(math.ceil(total_seconds)) + " Seconds")

    for url in vulnerable_urls:
        print("Vulnerable URL in " + url)

